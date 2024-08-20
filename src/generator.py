import docx
import re

# function to create a new invoice from a docx template
def generate_invoice(var_dict):
    # print("\n\n---TEMPLATE---\n")
    invoice = docx.Document("template.docx")
    # print all text in the template
    # for table in template.tables:
    #     for cell in table._cells:
            # print(cell.text)
    # replace placeholders
    # print("\n\n---REPLACEMENTS---\n\n")
    
    # replace all variables in tables
    for table in invoice.tables:
        for cell in table._cells:
            cell.text = find_and_replace(cell.text, var_dict)
    # replace vars in paragraphs
    for paragraph in invoice.paragraphs:
        paragraph.text = find_and_replace(paragraph.text, var_dict)
    # print("---FINISHED INVOICE---\n\n")
    # print all text in new invoice
    # for table in invoice.tables:
    #     for cell in table._cells:
    #         print(cell.text)
    print(f"Saving Invoice {var_dict['invoice_number']}... ")
    invoice.save(f"invoices/Invoice {var_dict['invoice_number']}.docx")
    print("Done!")
    

def find_and_replace(text, var_dict):
    # print("TEXT: ", text)
    matches = re.findall(r"placeholder_\w+", text)
    # print(matches)    
    vars = [match.replace("placeholder_", "") for match in matches]
    # print(vars)
    if matches:
        for i in range(len(matches)):
            # print(f"{matches[i]} -> {var_dict[vars[i]]}")
            text = text.replace(matches[i], var_dict[vars[i]])
    return text

def main():
    generate_invoice()

if __name__ == "__main__":
   main()